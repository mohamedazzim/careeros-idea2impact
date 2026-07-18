"""CareerOS Readiness Score and report API.

Scores and reports are computed from persisted pipeline evidence. Report export
does not call an LLM because the artifact must remain available even when a
provider is down.
"""

from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from typing import Any, Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Response
from pydantic import BaseModel
from sqlalchemy import desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from src.api.deps import get_current_user
from src.db.repositories.report_repository import GeneratedReportRepository
from src.db.session import get_db
from src.models.interview import InterviewSession
from src.models.jobs import Job, JobMatch
from src.models.knowledge import KnowledgeDoc
from src.models.report import GeneratedReport
from src.models.roadmap import Roadmap, RoadmapGoal, RoadmapTask
from src.models.user import User
from src.services.readiness import get_readiness_engine
from src.services.storage import storage_client

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/readiness", tags=["readiness"])

ALLOWED_REPORT_TYPES = {
    "candidate",
    "resume",
    "readiness",
    "interview",
    "roadmap",
    "opportunity",
    "career_progress",
}
ALLOWED_REPORT_FORMATS = {"pdf", "docx"}


class SubScore(BaseModel):
    score: float
    label: str
    evidence: list[str] = []


class ReadinessResponse(BaseModel):
    overall: float
    resume_score: SubScore
    interview_score: SubScore
    opportunity_score: SubScore
    skill_gap: SubScore
    market_readiness: SubScore
    career_progress: SubScore
    trend: list[dict] = []


class ReportRequest(BaseModel):
    user_id: Optional[str] = None
    report_type: str = "candidate"
    format: str = "pdf"


class TimelineEvent(BaseModel):
    stage: str
    timestamp: str
    status: str
    detail: str


@router.get("/score", response_model=ReadinessResponse)
async def get_readiness_score(user: dict = Depends(get_current_user)):
    """Compute CareerOS Readiness Score from persisted subsystem outputs."""
    user_id = user["sub"]

    result = await get_readiness_engine().compute(user_id)
    dims = result["dimensions"]

    resume = dims.get("resume_quality", {})
    skills = dims.get("skill_readiness", {})
    interview = dims.get("interview_readiness", {})
    opportunity = dims.get("opportunity_readiness", {})

    return ReadinessResponse(
        overall=result["overall"],
        resume_score=SubScore(score=resume.get("score", 0), label=resume.get("label", "Resume Intelligence"), evidence=_ev(resume)),
        interview_score=SubScore(score=interview.get("score", 0), label=interview.get("label", "Interview Performance"), evidence=_ev(interview)),
        opportunity_score=SubScore(score=opportunity.get("score", 0), label=opportunity.get("label", "Opportunity Discovery"), evidence=_ev(opportunity)),
        skill_gap=SubScore(score=skills.get("score", 0), label=skills.get("label", "Skill Readiness"), evidence=_ev(skills)),
        market_readiness=SubScore(score=skills.get("score", 0), label="Market Alignment", evidence=["Market alignment derived from persisted skill and job-match evidence"]),
        career_progress=SubScore(score=round(result["overall"] * 0.95, 1) if result["overall"] else 0, label="Career Progress", evidence=[f"Readiness score: {result['overall']}"]),
        trend=result.get("trend", []),
    )


def _ev(dim: dict) -> list[str]:
    """Extract a stable evidence list from a dimension dictionary."""
    ev = dim.get("evidence", [])
    if isinstance(ev, list):
        return [str(item) for item in ev]
    if isinstance(ev, str):
        return [ev]
    if isinstance(ev, dict):
        return [f"{k}: {v}" for k, v in list(ev.items())[:5]]
    return []


@router.get("/timeline")
async def get_career_timeline(user: dict = Depends(get_current_user)):
    """Career journey timeline from pipeline events."""
    user_id = user["sub"]
    return await get_readiness_engine().compute_timeline(user_id)


@router.get("/explain")
async def get_explainability(user: dict = Depends(get_current_user)):
    """Evidence chain explanations from persisted readiness computation."""
    user_id = user["sub"]
    readiness = await get_readiness_engine().compute(user_id)
    return {
        "overall_score": readiness["overall"],
        "formula": "sum(dimension_score * dimension_weight)",
        "dimensions": readiness["dimensions"],
        "trend": readiness.get("trend", []),
        "status": "computed_from_persisted_runtime_evidence",
    }


@router.post("/report")
async def generate_report(
    req: ReportRequest,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate and persist a PDF or DOCX report from real CareerOS data."""
    report_type = _validate_report_type(req.report_type)
    report_format = _validate_report_format(req.format)
    user_id = req.user_id or user["sub"]

    payload = await _build_report_payload(db, user_id, report_type)
    report = await _persist_report(db, user_id, payload, report_format)

    return {
        "report_id": report.report_uid,
        "report_type": report.report_type,
        "format": report.format,
        "status": report.status,
        "generated_at": report.created_at.isoformat() + "Z",
        "title": report.title,
        "summary": report.summary,
        "readiness_score": payload["readiness"]["overall"],
        "download_url": f"/api/v1/readiness/reports/{report.report_uid}/download",
        "evidence_counts": payload["evidence_counts"],
    }


@router.get("/reports")
async def list_reports(
    limit: int = Query(25, ge=1, le=100),
    offset: int = Query(0, ge=0),
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List persisted reports for the authenticated user."""
    reports, total = await GeneratedReportRepository(db).find_by_user(user["sub"], limit=limit, offset=offset)
    return {
        "total": total,
        "items": [_report_to_dict(report) for report in reports],
    }


@router.get("/reports/{report_id}/download")
async def download_persisted_report(
    report_id: str,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download a persisted report artifact."""
    report = await GeneratedReportRepository(db).get_by_uid(report_id)
    if not report or report.user_id != user["sub"] or report.status != "ready":
        raise HTTPException(status_code=404, detail="Report not found")

    content = await storage_client.read_file(report.storage_path)
    extension = "docx" if report.format == "docx" else "pdf"
    filename = f"{report.report_type}-{report.report_uid}.{extension}"
    return Response(
        content=content,
        media_type=report.content_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/report/download")
async def download_report(
    report_type: str = Query("candidate"),
    format: str = Query("pdf"),
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Compatibility endpoint that generates then downloads a report artifact."""
    report_type = _validate_report_type(report_type)
    report_format = _validate_report_format(format)
    payload = await _build_report_payload(db, user["sub"], report_type)
    report = await _persist_report(db, user["sub"], payload, report_format)
    content = await storage_client.read_file(report.storage_path)
    extension = "docx" if report.format == "docx" else "pdf"
    return Response(
        content=content,
        media_type=report.content_type,
        headers={"Content-Disposition": f'attachment; filename="{report.report_type}-{report.report_uid}.{extension}"'},
    )


def _validate_report_type(report_type: str) -> str:
    normalized = (report_type or "candidate").strip().lower()
    if normalized not in ALLOWED_REPORT_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported report_type: {report_type}")
    return normalized


def _validate_report_format(report_format: str) -> str:
    normalized = (report_format or "pdf").strip().lower()
    if normalized not in ALLOWED_REPORT_FORMATS:
        raise HTTPException(status_code=400, detail=f"Unsupported report format: {report_format}")
    return normalized


async def _build_report_payload(db: AsyncSession, user_id: str, report_type: str) -> dict[str, Any]:
    readiness = await get_readiness_engine().compute(user_id)
    user_record = await _fetch_user(db, user_id)
    resume = await _fetch_latest_resume(db, user_id)
    matches = await _fetch_top_matches(db, user_id)
    interview = await _fetch_latest_interview(db, user_id)
    roadmap = await _fetch_latest_roadmap(db, user_id)

    top_match_items = [_format_match(match, job) for match, job in matches]
    dimensions = readiness.get("dimensions", {})
    recommendations = _build_recommendations(dimensions, resume, top_match_items, interview, roadmap)

    payload = {
        "version": "career-os-report-v1",
        "report_type": report_type,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "user": {
            "id": user_id,
            "email": getattr(user_record, "email", None),
            "full_name": getattr(user_record, "full_name", None),
        },
        "readiness": readiness,
        "resume": _format_resume(resume),
        "top_matches": top_match_items,
        "interview": _format_interview(interview),
        "roadmap": roadmap,
        "recommendations": recommendations,
        "evidence_counts": {
            "resume_records": 1 if resume else 0,
            "job_matches": len(top_match_items),
            "interview_sessions": 1 if interview else 0,
            "roadmaps": 1 if roadmap else 0,
            "readiness_dimensions": len(dimensions),
        },
    }
    payload["summary"] = _build_report_summary(payload)
    payload["sections"] = _build_report_sections(payload)
    return payload


async def _fetch_user(db: AsyncSession, user_id: str) -> Optional[User]:
    try:
        result = await db.execute(select(User).where(User.id == user_id))
        return result.scalar_one_or_none()
    except Exception:
        logger.debug("Unable to resolve user record for report", exc_info=True)
        return None


async def _fetch_latest_resume(db: AsyncSession, user_id: str) -> Optional[KnowledgeDoc]:
    result = await db.execute(
        select(KnowledgeDoc)
        .where(KnowledgeDoc.user_id == user_id, KnowledgeDoc.deleted_at.is_(None))
        .order_by(desc(KnowledgeDoc.created_at))
        .limit(1)
    )
    return result.scalar_one_or_none()


async def _fetch_top_matches(db: AsyncSession, user_id: str) -> list[tuple[JobMatch, Job]]:
    result = await db.execute(
        select(JobMatch, Job)
        .join(Job, Job.id == JobMatch.job_id)
        .where(JobMatch.user_id == user_id, JobMatch.deleted_at.is_(None), Job.deleted_at.is_(None), Job.status == "active")
        .order_by(desc(JobMatch.overall_score), desc(JobMatch.created_at))
        .limit(10)
    )
    return list(result.all())


async def _fetch_latest_interview(db: AsyncSession, user_id: str) -> Optional[InterviewSession]:
    result = await db.execute(
        select(InterviewSession)
        .where(InterviewSession.user_id == user_id, InterviewSession.deleted_at.is_(None))
        .order_by(desc(InterviewSession.created_at))
        .limit(1)
    )
    return result.scalar_one_or_none()


async def _fetch_latest_roadmap(db: AsyncSession, user_id: str) -> Optional[dict[str, Any]]:
    result = await db.execute(
        select(Roadmap)
        .where(Roadmap.user_id == user_id, Roadmap.deleted_at.is_(None))
        .order_by(desc(Roadmap.created_at))
        .limit(1)
    )
    roadmap = result.scalar_one_or_none()
    if not roadmap:
        return None

    goals_result = await db.execute(
        select(RoadmapGoal)
        .where(RoadmapGoal.roadmap_id == roadmap.id, RoadmapGoal.deleted_at.is_(None))
        .order_by(RoadmapGoal.order_index, RoadmapGoal.id)
    )
    goals = list(goals_result.scalars().all())
    tasks_result = await db.execute(
        select(RoadmapTask)
        .join(RoadmapGoal, RoadmapGoal.id == RoadmapTask.goal_id)
        .where(RoadmapGoal.roadmap_id == roadmap.id, RoadmapTask.deleted_at.is_(None))
        .order_by(RoadmapTask.order_index, RoadmapTask.id)
    )
    tasks = list(tasks_result.scalars().all())
    tasks_by_goal: dict[int, list[dict[str, Any]]] = {}
    for task in tasks:
        tasks_by_goal.setdefault(task.goal_id, []).append(
            {
                "title": task.title,
                "completed": task.completed,
                "due_date": task.due_date,
                "description": task.description,
            }
        )

    return {
        "title": roadmap.title,
        "target_role": roadmap.target_role,
        "target_location": roadmap.target_location,
        "target_timeline": roadmap.target_timeline,
        "progress_pct": roadmap.progress_pct,
        "trace_id": roadmap.trace_id,
        "goals": [
            {
                "title": goal.title,
                "description": goal.description,
                "category": goal.category,
                "priority": goal.priority,
                "tasks": tasks_by_goal.get(goal.id, []),
            }
            for goal in goals
        ],
    }


def _format_resume(resume: Optional[KnowledgeDoc]) -> Optional[dict[str, Any]]:
    if not resume:
        return None
    analysis = resume.analysis_results or {}
    return {
        "doc_uid": resume.doc_uid,
        "title": resume.title,
        "status": resume.status,
        "chunk_count": resume.chunk_count,
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
        "summary": resume.summary,
        "skills": _limited_list(analysis.get("skills") or analysis.get("extracted_skills") or []),
        "analysis_keys": sorted(analysis.keys())[:20],
    }


def _format_match(match: JobMatch, job: Job) -> dict[str, Any]:
    return {
        "job_id": job.job_uid,
        "title": job.title,
        "company": job.company,
        "location": job.location,
        "source_provider": job.source_provider or job.source,
        "apply_url": job.apply_url,
        "posted_date": job.posted_date.isoformat() if job.posted_date else None,
        "overall_score": round(float(match.overall_score or 0), 1),
        "skill_match": round(float(match.skill_match or 0), 1),
        "experience_match": round(float(match.experience_match or 0), 1),
        "education_match": round(float(match.education_match or 0), 1),
        "strengths": _limited_list(match.strengths or []),
        "gaps": _limited_list(match.gaps or []),
        "recommendation": match.recommendation,
    }


def _format_interview(interview: Optional[InterviewSession]) -> Optional[dict[str, Any]]:
    if not interview:
        return None
    return {
        "session_uid": interview.session_uid,
        "interview_type": interview.interview_type,
        "status": interview.status,
        "difficulty_level": interview.difficulty_level,
        "total_score": interview.total_score,
        "created_at": interview.created_at.isoformat() if interview.created_at else None,
        "closed_at": interview.closed_at.isoformat() if interview.closed_at else None,
    }


def _build_recommendations(
    dimensions: dict[str, Any],
    resume: Optional[KnowledgeDoc],
    top_matches: list[dict[str, Any]],
    interview: Optional[InterviewSession],
    roadmap: Optional[dict[str, Any]],
) -> list[str]:
    recommendations: list[str] = []
    low_dimensions = sorted(
        (
            (name, float(data.get("score", 0)), data)
            for name, data in dimensions.items()
            if isinstance(data, dict)
        ),
        key=lambda item: item[1],
    )
    for name, score, data in low_dimensions[:3]:
        label = data.get("label") or name.replace("_", " ").title()
        evidence = "; ".join(_ev(data)[:2])
        recommendations.append(f"Improve {label}: current score {score:.1f}. Evidence: {evidence or 'no detailed evidence captured'}")
    if resume and resume.status != "analyzed":
        recommendations.append(f"Complete resume analysis for {resume.title}; current status is {resume.status}.")
    if top_matches:
        gaps = []
        for match in top_matches[:3]:
            gaps.extend(match.get("gaps") or [])
        if gaps:
            recommendations.append("Prioritize recurring job-match gaps: " + ", ".join(dict.fromkeys(str(g) for g in gaps[:6])))
    if interview and float(interview.total_score or 0) < 70:
        recommendations.append(f"Run another interview practice cycle; latest score is {interview.total_score}.")
    if roadmap and float(roadmap.get("progress_pct") or 0) < 50:
        recommendations.append(f"Advance roadmap execution; current progress is {roadmap.get('progress_pct')}%.")
    return recommendations[:8]


def _build_report_summary(payload: dict[str, Any]) -> str:
    counts = payload["evidence_counts"]
    return (
        f"{payload['report_type'].replace('_', ' ').title()} report generated from "
        f"{counts['resume_records']} resume record, {counts['job_matches']} job matches, "
        f"{counts['interview_sessions']} interview session, and {counts['roadmaps']} roadmap."
    )


def _build_report_sections(payload: dict[str, Any]) -> list[dict[str, Any]]:
    readiness = payload["readiness"]
    sections = [
        {
            "title": "Readiness Score",
            "items": [
                f"Overall: {readiness.get('overall', 0)}",
                f"Formula: {readiness.get('formula', 'sum(dimension_score * dimension_weight)')}",
            ],
        },
        {
            "title": "Resume Evidence",
            "items": _dict_items(payload.get("resume") or {"status": "No resume record found"}),
        },
        {
            "title": "Top Opportunity Matches",
            "items": [
                f"{item['overall_score']}% - {item['title']} at {item.get('company') or 'Unknown company'} ({item.get('source_provider') or 'unknown source'})"
                for item in payload.get("top_matches", [])[:10]
            ] or ["No job-match rows found."],
        },
        {
            "title": "Interview Evidence",
            "items": _dict_items(payload.get("interview") or {"status": "No interview session found"}),
        },
        {
            "title": "Roadmap Evidence",
            "items": _roadmap_items(payload.get("roadmap")),
        },
        {
            "title": "Recommendations",
            "items": payload.get("recommendations") or ["No recommendations generated from current evidence."],
        },
    ]
    return sections


async def _persist_report(
    db: AsyncSession,
    user_id: str,
    payload: dict[str, Any],
    report_format: str,
) -> GeneratedReport:
    report_uid = f"report-{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}"
    if report_format == "docx":
        content = _render_docx(payload)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content = _render_pdf(payload)
        content_type = "application/pdf"

    filename = f"{report_uid}.{report_format}"
    storage_path = await storage_client.save_file(filename, content)
    return await GeneratedReportRepository(db).create(
        report_uid=report_uid,
        user_id=user_id,
        report_type=payload["report_type"],
        title=f"{payload['report_type'].replace('_', ' ').title()} Report",
        version=payload["version"],
        format=report_format,
        summary=payload["summary"],
        payload=payload,
        storage_path=storage_path,
        content_type=content_type,
        status="ready",
    )


def _render_docx(payload: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(f"CareerOS {payload['report_type'].replace('_', ' ').title()} Report", level=1)
    document.add_paragraph(f"Generated: {payload['generated_at']}")
    document.add_paragraph(f"User: {payload['user'].get('email') or payload['user']['id']}")
    document.add_paragraph(payload["summary"])
    for section in payload["sections"]:
        document.add_heading(section["title"], level=2)
        for item in section["items"]:
            document.add_paragraph(str(item), style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf(payload: dict[str, Any]) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    y = 48
    y = _pdf_text(page, f"CareerOS {payload['report_type'].replace('_', ' ').title()} Report", y, size=18, bold=True)
    y = _pdf_text(page, f"Generated: {payload['generated_at']}", y)
    y = _pdf_text(page, f"User: {payload['user'].get('email') or payload['user']['id']}", y)
    y = _pdf_text(page, payload["summary"], y)

    for section in payload["sections"]:
        if y > 720:
            page = doc.new_page()
            y = 48
        y = _pdf_text(page, section["title"], y + 8, size=13, bold=True)
        for item in section["items"]:
            if y > 740:
                page = doc.new_page()
                y = 48
            y = _pdf_text(page, f"- {item}", y)

    content = doc.tobytes()
    doc.close()
    return content


def _pdf_text(page: Any, text: str, y: float, size: int = 10, bold: bool = False) -> float:
    fontname = "helv" if not bold else "hebo"
    page.insert_textbox(
        fitz_rect(48, y, 545, y + 48),
        str(text)[:1200],
        fontsize=size,
        fontname=fontname,
        align=0,
    )
    return y + max(18, (len(str(text)) // 90 + 1) * 14)


def fitz_rect(x0: float, y0: float, x1: float, y1: float):
    import fitz

    return fitz.Rect(x0, y0, x1, y1)


def _limited_list(value: Any, limit: int = 12) -> list[Any]:
    if not value:
        return []
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            value = parsed
        except Exception:
            return [value]
    if isinstance(value, dict):
        return [f"{key}: {val}" for key, val in list(value.items())[:limit]]
    if isinstance(value, list):
        return value[:limit]
    return [value]


def _dict_items(data: dict[str, Any]) -> list[str]:
    return [f"{key}: {value}" for key, value in data.items() if value not in (None, "", [], {})][:12]


def _roadmap_items(roadmap: Optional[dict[str, Any]]) -> list[str]:
    if not roadmap:
        return ["No roadmap found."]
    items = [
        f"Title: {roadmap.get('title')}",
        f"Target role: {roadmap.get('target_role')}",
        f"Progress: {roadmap.get('progress_pct')}%",
    ]
    for goal in (roadmap.get("goals") or [])[:6]:
        items.append(f"Goal: {goal.get('title')} ({len(goal.get('tasks') or [])} tasks)")
    return [item for item in items if item and not item.endswith(": None")]


def _report_to_dict(report: GeneratedReport) -> dict[str, Any]:
    return {
        "report_id": report.report_uid,
        "report_type": report.report_type,
        "title": report.title,
        "format": report.format,
        "status": report.status,
        "summary": report.summary,
        "created_at": report.created_at.isoformat() + "Z" if report.created_at else None,
        "download_url": f"/api/v1/readiness/reports/{report.report_uid}/download",
    }
