"""
Production-grade document parsing pipeline.
Async-safe, worker-safe, retry-safe extraction from PDF, DOCX, ODT, RTF, TXT, and images.
Malformed PDF handling, image-based PDF detection, section structure preservation.
Stateless — ready for LangGraph node extraction.
"""
import asyncio
import io
import logging
import re
import time
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, Any, Optional, List, Tuple

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError as DocxPackageError

from src.observability.metrics import (
    PARSER_COUNT,
    PARSER_LATENCY,
    PARSER_BYTES,
    PARSER_PAGE_COUNT,
)
from src.observability.langsmith import traceable
from .interfaces import (
    ParseResult,
    PermanentPipelineError,
    RetryablePipelineError,
    ProcessingStatus,
)
from .mime_detector import mime_detector

logger = logging.getLogger(__name__)

# Shared thread pool for CPU-bound parse operations, max 4 workers
_PARSE_EXECUTOR = ThreadPoolExecutor(max_workers=4, thread_name_prefix="parser-")

# Section detection patterns
SECTION_HEADER_RE = re.compile(
    r'^(?:EXPERIENCE|EDUCATION|SKILLS?|EMPLOYMENT|QUALIFICATIONS|'
    r'PROJECTS?|CERTIFICATIONS?|SUMMARY|PROFILE|OBJECTIVE|'
    r'ACHIEVEMENTS?|LANGUAGES?|INTERESTS?|REFERENCES?|'
    r'WORK\s+HISTORY|PROFESSIONAL\s+EXPERIENCE|TECHNICAL\s+SKILLS?|'
    r'CONTACT|PERSONAL\s+INFORMATION)',
    re.IGNORECASE | re.MULTILINE
)

# Minimum text length to consider a page as having extractable text
MIN_PAGE_TEXT_LENGTH = 30

# MIME types that require OCR (no direct text extraction possible)
IMAGE_MIME_TYPES = frozenset({
    "image/png", "image/jpeg", "image/gif", "image/bmp", "image/tiff",
    "image/jpg",
})


class InsufficientContentError(PermanentPipelineError):
    pass


class MalformedDocumentError(RetryablePipelineError):
    """Document is structurally malformed but may be recoverable."""
    pass


class ParserPipeline:
    """Production-grade document parsing pipeline. Fully async-safe via thread pool."""

    MIN_CONTENT_LENGTH = 50
    MIN_WORD_COUNT = 5

    async def parse(
        self,
        filename: str,
        storage_path: str,
        content_type: Optional[str] = None,
    ) -> ParseResult:
        """
        Parse document to extract text with metadata.
        Routes to the appropriate format-specific parser.

        Returns:
            ParseResult with extracted text, content_type, metadata, confidence

        Raises:
            FileNotFoundError: If source file missing
            InsufficientContentError: If content too short
            RetryablePipelineError: For transient/malformed failures
        """
        logger.info(f"Parsing document: {filename}", extra={
            "filename": filename,
            "storage_path": storage_path,
            "content_type": content_type,
        })

        start = time.monotonic()

        # Determine content type
        if not content_type:
            content_type = mime_detector.detect_from_extension(filename) or "application/octet-stream"

        # Read file bytes (for formats that need it)
        file_bytes = None
        try:
            with open(storage_path, "rb") as f:
                file_bytes = f.read()
        except FileNotFoundError:
            raise PermanentPipelineError(f"Source file not found: {storage_path}")

        if file_bytes:
            PARSER_BYTES.labels(format=self._format_label(content_type)).observe(len(file_bytes))

        # Route to format-specific parser
        try:
            if content_type == "application/pdf":
                text, metadata = await self._parse_pdf(storage_path, file_bytes)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text, metadata = await self._parse_docx(storage_path, file_bytes)
            elif content_type == "application/vnd.oasis.opendocument.text":
                text, metadata = await self._parse_odt(file_bytes)
            elif content_type == "application/rtf":
                text, metadata = await self._parse_rtf(file_bytes)
            elif content_type in ("text/plain", "text/html"):
                text, metadata = await self._parse_text(file_bytes)
            elif mime_detector.is_image(content_type):
                text, metadata = await self._parse_image(storage_path, file_bytes, filename)
            elif content_type == "application/msword":
                text, metadata = await self._parse_doc(file_bytes)
            else:
                # Fallback: try as text
                text, metadata = await self._parse_text(file_bytes)
                content_type = "text/plain"
        except PermanentPipelineError:
            raise
        except Exception as e:
            logger.error(f"Document parsing failed: {e}", extra={
                "doc_filename": filename,
                "error_details": str(e),
            })
            PARSER_COUNT.labels(format=self._format_label(content_type), status="error").inc()
            raise RetryablePipelineError(f"Document parsing failed: {e}")

        elapsed = time.monotonic() - start
        PARSER_LATENCY.labels(format=self._format_label(content_type)).observe(elapsed)

        # Validate content — skip for image types that require OCR
        stripped = text.strip()
        word_count = len(stripped.split()) if stripped else 0

        is_image = content_type in IMAGE_MIME_TYPES or (
            metadata.get("extraction_method") == "requires_ocr"
        )

        if not is_image and word_count < self.MIN_WORD_COUNT:
            PARSER_COUNT.labels(format=self._format_label(content_type), status="insufficient").inc()
            raise InsufficientContentError(
                f"Document parsing yielded insufficient content "
                f"(words: {word_count}, length: {len(stripped)})"
            )

        PARSER_COUNT.labels(format=self._format_label(content_type), status="success").inc()

        if metadata.get("page_count"):
            PARSER_PAGE_COUNT.labels(format=self._format_label(content_type)).observe(
                metadata["page_count"]
            )

        # Detect sections in the text
        sections = metadata.get("sections", [])
        if not sections:
            sections = self._detect_sections(text)

        result = ParseResult(
            text=text,
            content_type=content_type,
            metadata={
                "filename": filename,
                "original_length": len(text),
                "stripped_length": len(stripped),
                "word_count": word_count,
                "page_count": metadata.get("page_count", 1),
                "parse_duration_ms": round(elapsed * 1000, 2),
                "sections": sections,
                "image_based_pages": metadata.get("image_based_pages", []),
                "extraction_method": metadata.get("extraction_method", "standard"),
                "document_metadata": metadata.get("document_metadata", {}),
            },
            confidence=metadata.get("confidence", 1.0),
        )

        logger.info("Document parsed successfully", extra={
            "filename": filename,
            "content_length": len(stripped),
            "word_count": word_count,
            "page_count": metadata.get("page_count", 1),
        })

        return result

    # ── PDF Parsing ──────────────────────────────────────────────────

    @traceable(name="parse_pdf", run_type="chain")
    async def _parse_pdf(
        self, storage_path: str, file_bytes: bytes
    ) -> Tuple[str, dict]:
        """Parse PDF with malformed file handling and image-based page detection."""
        return await asyncio.get_event_loop().run_in_executor(
            _PARSE_EXECUTOR, self._parse_pdf_sync, storage_path, file_bytes
        )

    def _parse_pdf_sync(self, storage_path: str, file_bytes: bytes) -> Tuple[str, dict]:
        """Synchronous PDF extraction (runs in thread pool)."""
        text_parts = []
        image_based_pages = []
        page_count = 0
        doc_metadata = {}
        had_error = False

        doc = None
        try:
            doc = fitz.open(storage_path)
        except Exception:
            # Try stream-based open for malformed on-disk PDF
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
            except Exception as e:
                raise MalformedDocumentError(f"Failed to open PDF: {e}")

        try:
            page_count = len(doc)
            doc_metadata = dict(doc.metadata) if doc.metadata else {}

            for i, page in enumerate(doc):
                page_text = ""
                extraction_method = "pdf_text"

                # Try native text extraction first
                try:
                    page_text = page.get_text("text").strip()
                except Exception:
                    pass

                # Detect image-based pages (very little or no extractable text)
                if len(page_text) < MIN_PAGE_TEXT_LENGTH:
                    # Try text blocks extraction (more granular)
                    try:
                        blocks = page.get_text("blocks")
                        page_text = "\n".join(
                            b[4].strip() for b in blocks
                            if len(b) >= 5 and isinstance(b[4], str) and b[4].strip()
                        )
                    except Exception:
                        page_text = ""

                # If still insufficient text, mark as image-based
                if len(page_text) < MIN_PAGE_TEXT_LENGTH:
                    # Try structured dict extraction
                    try:
                        text_dict = page.get_text("dict")
                        extracted = []
                        for block in text_dict.get("blocks", []):
                            if block.get("type") == 0:  # text block
                                for line in block.get("lines", []):
                                    line_spans = [
                                        span.get("text", "")
                                        for span in line.get("spans", [])
                                    ]
                                    extracted.append("".join(line_spans))
                        page_text = "\n".join(e for e in extracted if e.strip())
                    except Exception:
                        pass

                if len(page_text) < MIN_PAGE_TEXT_LENGTH:
                    image_based_pages.append(i)
                    extraction_method = "ocr"
                    page_text = "[IMAGE-BASED PAGE — requires OCR]"

                # Preserve page boundaries
                text_parts.append(page_text)

                if i == 0:
                    doc_metadata["page0_method"] = extraction_method
        finally:
            if doc:
                try:
                    doc.close()
                except Exception:
                    pass

        text = "\n\n".join(text_parts)

        # Detect sections from font-size heuristics (from first few pages)
        sections = self._extract_pdf_sections(doc, page_count) if not had_error else []

        metadata = {
            "page_count": page_count,
            "image_based_pages": image_based_pages,
            "has_image_pages": len(image_based_pages) > 0,
            "image_page_ratio": len(image_based_pages) / max(page_count, 1),
            "extraction_method": "standard" if not image_based_pages else "mixed",
            "document_metadata": doc_metadata,
            "sections": sections,
            "confidence": 0.95 if not image_based_pages else 0.95 * (1 - len(image_based_pages) / page_count),
        }

        return text, metadata

    def _extract_pdf_sections(self, doc, page_count: int) -> List[str]:
        """Extract section-like headings from PDF using font size heuristics."""
        sections = []
        try:
            for i in range(min(page_count, 3)):  # First 3 pages only
                page = doc[i]
                blocks = page.get_text("dict").get("blocks", [])
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            size = span.get("size", 0)
                            is_bold = "Bold" in (span.get("font", "") or "")
                            if text and size >= 13 and (is_bold or SECTION_HEADER_RE.match(text)):
                                if text not in sections:
                                    sections.append(text)
        except Exception:
            pass
        return sections

    # ── DOCX Parsing ─────────────────────────────────────────────────

    @traceable(name="parse_docx", run_type="chain")
    async def _parse_docx(
        self, storage_path: str, file_bytes: bytes
    ) -> Tuple[str, dict]:
        """Parse DOCX with style hierarchy preservation."""
        return await asyncio.get_event_loop().run_in_executor(
            _PARSE_EXECUTOR, self._parse_docx_sync, storage_path, file_bytes
        )

    def _parse_docx_sync(self, storage_path: str, file_bytes: bytes) -> Tuple[str, dict]:
        """Synchronous DOCX extraction with formatting hierarchy."""
        try:
            doc = DocxDocument(storage_path)
        except (DocxPackageError, Exception) as e:
            raise MalformedDocumentError(f"Failed to open DOCX: {e}")

        sections_found = []
        full_text_lines = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect heading level from style
            style_name = para.style.name.lower() if para.style else ""
            is_heading = style_name.startswith("heading") or style_name.startswith("title")

            if is_heading:
                # Section headers get special treatment
                sections_found.append(text)
                current_section = text
                full_text_lines.append(f"\n{text.upper()}\n")
            elif current_section:
                full_text_lines.append(text)
            else:
                full_text_lines.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    full_text_lines.append(row_text)

        # Extract core properties as metadata
        doc_metadata = {}
        core = doc.core_properties
        for attr in ("title", "author", "subject", "category", "keywords"):
            val = getattr(core, attr, None)
            if val:
                doc_metadata[attr] = str(val)

        text = "\n".join(full_text_lines)

        # Detect sections from the full text if style-based detection found nothing
        if not sections_found:
            sections_found = self._detect_sections(text)

        metadata = {
            "page_count": 1,
            "sections": sections_found,
            "document_metadata": doc_metadata,
            "extraction_method": "docx_structured",
            "confidence": 0.98,
        }

        return text, metadata

    # ── ODT Parsing ──────────────────────────────────────────────────

    async def _parse_odt(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Parse ODT via ZIP content.xml extraction."""
        return await asyncio.get_event_loop().run_in_executor(
            _PARSE_EXECUTOR, self._parse_odt_sync, file_bytes
        )

    def _parse_odt_sync(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Extract text from ODT content.xml."""
        import zipfile
        from xml.etree import ElementTree

        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                with zf.open("content.xml") as f:
                    tree = ElementTree.parse(f)

            ns = {
                "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
                "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
            }
            paragraphs = tree.findall(".//text:p", ns)
            text_parts = []
            for p in paragraphs:
                line = "".join(p.itertext()).strip()
                if line:
                    text_parts.append(line)

            text = "\n".join(text_parts)
            sections = self._detect_sections(text)

            return text, {
                "page_count": 1,
                "sections": sections,
                "extraction_method": "odt_xml",
                "confidence": 0.95,
            }
        except Exception as e:
            raise RetryablePipelineError(f"ODT parsing failed: {e}")

    # ── RTF Parsing ──────────────────────────────────────────────────

    async def _parse_rtf(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Basic RTF extraction — strips RTF control codes."""
        return await asyncio.get_event_loop().run_in_executor(
            _PARSE_EXECUTOR, self._parse_rtf_sync, file_bytes
        )

    def _parse_rtf_sync(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Strip RTF markup and extract plain text."""
        text = file_bytes.decode("ascii", errors="ignore")

        # Strip RTF grouping braces
        text = re.sub(r'\\\{', '{', text)
        text = re.sub(r'\\\}', '}', text)
        text = re.sub(r'\\\\(?![a-z])', '\\\\', text)

        # Remove control words
        text = re.sub(r'\\([a-z]+)(-?\d+)?[ ]?', '', text, flags=re.IGNORECASE)
        # Remove escaped characters
        text = re.sub(r"\\'([0-9a-fA-F]{2})", '', text)
        # Remove remaining braces
        text = text.replace('{', '').replace('}', '')
        # Normalize whitespace
        text = re.sub(r'\\par\b', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\\line\b', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\\tab\b', '\t', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*\n', '\n\n', text)

        lines = [l.strip() for l in text.splitlines() if l.strip()]
        text = "\n".join(lines)
        sections = self._detect_sections(text)

        return text, {
            "page_count": 1,
            "sections": sections,
            "extraction_method": "rtf_stripped",
            "confidence": 0.85,
        }

    # ── Legacy .doc Parsing ──────────────────────────────────────────

    async def _parse_doc(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Attempt to extract text from legacy .doc (binary) format."""
        return await asyncio.get_event_loop().run_in_executor(
            _PARSE_EXECUTOR, self._parse_doc_sync, file_bytes
        )

    def _parse_doc_sync(self, file_bytes: bytes) -> Tuple[str, dict]:
        """
        Extract readable ASCII/UTF-8 segments from binary .doc.
        This is a best-effort fallback — .doc is a binary format.
        """
        text = file_bytes.decode("ascii", errors="ignore")
        # Extract printable segments of reasonable length
        lines = []
        current = []
        for ch in text:
            if ch.isprintable() or ch in '\n\r\t ':
                current.append(ch)
            else:
                if len(current) >= 4:
                    segment = "".join(current).strip()
                    if len(segment) > 5:
                        lines.append(segment)
                current = []
        if len(current) >= 4:
            segment = "".join(current).strip()
            if len(segment) > 5:
                lines.append(segment)

        # Deduplicate consecutive empty lines
        cleaned = []
        prev_empty = False
        for line in lines:
            is_empty = not line.strip()
            if is_empty and prev_empty:
                continue
            prev_empty = is_empty
            cleaned.append(line)

        text = "\n".join(cleaned)
        sections = self._detect_sections(text)

        return text, {
            "page_count": 1,
            "sections": sections,
            "extraction_method": "doc_binary_fallback",
            "confidence": 0.60,
        }

    # ── Plain Text Parsing ───────────────────────────────────────────

    async def _parse_text(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Parse plain text or HTML files."""
        try:
            text = file_bytes.decode("utf-8", errors="strict")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="ignore")

        # Strip HTML tags if present
        if text.strip().startswith("<") and mime_detector._is_html(file_bytes):
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            text = re.sub(r'&[a-z]+;', ' ', text, flags=re.IGNORECASE)

        sections = self._detect_sections(text)

        return text, {
            "page_count": 1,
            "sections": sections,
            "extraction_method": "text_decode",
            "confidence": 1.0,
        }

    # ── Image Parsing (OCR-delegated) ────────────────────────────────

    async def _parse_image(
        self, storage_path: str, file_bytes: bytes, filename: str
    ) -> Tuple[str, dict]:
        """
        Image parsing — delegates to OCR.
        The actual OCR is handled by the OcrPipeline stage.
        Here we set up metadata and pass through.
        """
        return (
            "",
            {
                "page_count": 1,
                "image_based_pages": [0],
                "has_image_pages": True,
                "image_page_ratio": 1.0,
                "extraction_method": "requires_ocr",
                "confidence": 0.0,
                "filename": filename,
                "storage_path": storage_path,
            },
        )

    # ── Section Detection ────────────────────────────────────────────

    def _detect_sections(self, text: str) -> List[str]:
        """Detect resume section headers in text."""
        sections = []
        for line in text.splitlines():
            stripped = line.strip()
            if SECTION_HEADER_RE.match(stripped) and len(stripped) < 60:
                if stripped not in sections:
                    sections.append(stripped)
        return sections

    # ── Helpers ──────────────────────────────────────────────────────

    def _format_label(self, content_type: str) -> str:
        """Short label for metric dimensions."""
        return content_type.split("/")[-1].replace(".", "_") if "/" in content_type else content_type

    # ── LangGraph node interface ─────────────────────────────────────

    async def run(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """
        LangGraph node entry point.

        Args:
            state: ProcessingState dict

        Returns:
            State update dict
        """
        try:
            result = await self.parse(
                filename=state["filename"],
                storage_path=state["storage_path"],
                content_type=state.get("content_type"),
            )

            return {
                "raw_text": result.text,
                "content_type": result.content_type,
                "parse_error": None,
                "status": ProcessingStatus.EXTRACTING,
            }
        except PermanentPipelineError as e:
            return {
                "raw_text": None,
                "parse_error": str(e),
                "status": ProcessingStatus.FAILED,
            }
        except RetryablePipelineError as e:
            return {
                "raw_text": None,
                "parse_error": str(e),
                "status": ProcessingStatus.PARSING,
            }


parser_pipeline = ParserPipeline()
