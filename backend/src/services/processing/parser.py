import io
import logging
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from .ocr_service import ocr_service
from langsmith import traceable

logger = logging.getLogger(__name__)

class DocumentParser:
    @traceable(name="parse_document")
    def parse_document(self, file_bytes: bytes = None, filename: str = "", file_path: str = None) -> dict:
        metadata = {}
        text = ""
        document_type = "unknown"
        page_count = 0
        
        try:
            filename_lower = filename.lower()
            if filename_lower.endswith(".pdf"):
                text, page_count = self._parse_pdf(file_bytes=file_bytes, file_path=file_path)
                document_type = "pdf"
            elif filename_lower.endswith(".docx"):
                text, page_count = self._parse_docx(file_bytes=file_bytes, file_path=file_path)
                document_type = "docx"
            elif filename_lower.endswith((".txt", ".md")):
                if file_path:
                    with open(file_path, 'rb') as f:
                        file_bytes = f.read()
                text = file_bytes.decode('utf-8', errors='ignore')
                document_type = "txt"
                page_count = 1
            elif filename_lower.endswith((".png", ".jpg", ".jpeg")):
                if file_path:
                    with open(file_path, 'rb') as f:
                        file_bytes = f.read()
                text = ocr_service.extract_text_from_image(file_bytes)
                document_type = "image"
                page_count = 1
            else:
                # Try generic decode
                if file_path and not file_bytes:
                    with open(file_path, 'rb') as f:
                        file_bytes = f.read()
                text = file_bytes.decode('utf-8', errors='ignore')
                document_type = "txt"
                page_count = 1
                
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            raise ValueError(f"Failed to parse document: {str(e)}")

        return {
            "text": text,
            "metadata": metadata,
            "document_type": document_type,
            "page_count": page_count
        }

    def _parse_pdf(self, file_bytes: bytes = None, file_path: str = None) -> tuple[str, int]:
        text_parts = []
        if file_path:
            doc = fitz.open(file_path)
        else:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        
        for i, page in enumerate(doc):
            page_text = page.get_text("text").strip()
            if not page_text:
                # OCR Fallback
                try:
                    pix = page.get_pixmap()
                    img_bytes = pix.tobytes("png")
                    page_text = ocr_service.extract_text_from_image(img_bytes)
                except Exception as e:
                    logger.warning(f"OCR fallback failed on page {i}: {e}")
            
            text_parts.append(page_text)
            
        doc.close()
        return "\n\n".join(text_parts), page_count

    def _parse_docx(self, file_bytes: bytes = None, file_path: str = None) -> tuple[str, int]:
        if file_path:
            doc = DocxDocument(file_path)
        else:
            doc = DocxDocument(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # docx page count is non-trivial, store core metadata
        page_count = len(doc.core_properties.title or "") > 0 and 1 or 1 
        return "\n".join(full_text), page_count

document_parser = DocumentParser()
