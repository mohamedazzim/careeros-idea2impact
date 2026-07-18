"""
Extraction benchmarks for document parsing and OCR pipelines.
Measures accuracy, latency, throughput, and failure recovery.
"""
import os
import shutil
import time
import statistics
from pathlib import Path

import pytest
from PIL import Image, ImageDraw, ImageFont

from src.services.resume.processing.mime_detector import MimeDetectorPipeline
from src.services.resume.processing.parser import ParserPipeline, InsufficientContentError, PermanentPipelineError, RetryablePipelineError
from src.services.resume.processing.ocr_pipeline import OcrPipeline

TEST_FILES_DIR = Path(__file__).parent / "test_files"

requires_tesseract = pytest.mark.skipif(
    shutil.which("tesseract") is None and os.getenv("RUN_OCR_TESTS") != "true",
    reason="OCR integration test requires tesseract binary",
)


# ── Benchmark fixtures ──────────────────────────────────────────────

RESUME_TEXT = """John Alexander Smith
Software Engineer | San Francisco, CA
john.smith@example.com | (555) 123-4567

SUMMARY
Senior software engineer with 8+ years of experience building scalable web applications.
Proficient in Python, React, TypeScript, and cloud infrastructure.

EXPERIENCE
Senior Software Engineer, Acme Corp
January 2019 - Present
- Led migration of monolithic application to microservices architecture
- Improved API response time by 40% through caching and query optimization
- Mentored team of 5 junior engineers

Software Engineer, TechStart Inc.
June 2016 - December 2018
- Developed customer-facing dashboard used by 50,000+ users
- Implemented real-time notification system using WebSockets
- Reduced deployment time from 2 hours to 15 minutes via CI/CD

EDUCATION
Master of Science, Computer Science
Stanford University, 2014-2016

Bachelor of Science, Computer Science
UC Berkeley, 2010-2014

SKILLS
Programming: Python, JavaScript, TypeScript, Go, Java
Frontend: React, Vue.js, HTML5, CSS3, Tailwind
Backend: FastAPI, Django, Node.js, Express
Databases: PostgreSQL, MongoDB, Redis, Elasticsearch
DevOps: Docker, Kubernetes, AWS, Terraform, CI/CD
Tools: Git, JIRA, Confluence, Datadog

CERTIFICATIONS
AWS Solutions Architect Professional
Google Cloud Professional Data Engineer

LANGUAGES
English (Native), Spanish (Fluent), Mandarin (Basic)

PROJECTS
OpenSourceML: Contributed to popular Python ML library
DevTools: Built CLI productivity tool with 2,000+ GitHub stars
"""

BENCHMARK_SIZES = [
    ("small", RESUME_TEXT[:500]),
    ("medium", RESUME_TEXT[:1500]),
    ("large", RESUME_TEXT * 5),
]

CHARS_TO_EXTRACT = [
    ("noise_0", 0),
    ("noise_10", 10),
    ("noise_50", 50),
]


# ── MIME Detection Benchmarks ──────────────────────────────────────

class TestMimeBenchmarks:
    """Benchmark MIME detection speed and accuracy."""

    @pytest.fixture
    def detector(self):
        return MimeDetectorPipeline()

    def test_detection_speed_pdf(self, detector):
        """PDF detection should be near-instant."""
        content = b"%PDF-1.7\n...\n%%EOF"
        iterations = 1000
        start = time.perf_counter()
        for _ in range(iterations):
            detector.detect_from_content(content)
        elapsed = time.perf_counter() - start
        avg_us = (elapsed / iterations) * 1_000_000
        assert avg_us < 500, f"PDF detection too slow: {avg_us:.1f}µs per call"

    def test_detection_speed_image(self, detector):
        """Image magic number detection should be near-instant."""
        content = b'\x89PNG\r\n\x1a\n' + b'\x00' * 100
        iterations = 1000
        start = time.perf_counter()
        for _ in range(iterations):
            detector.detect_from_content(content)
        elapsed = time.perf_counter() - start
        avg_us = (elapsed / iterations) * 1_000_000
        assert avg_us < 500, f"Image detection too slow: {avg_us:.1f}µs per call"

    @pytest.mark.parametrize("label,text", BENCHMARK_SIZES)
    def test_mime_validation_speed(self, detector, label, text):
        """MIME validation should scale minimally with content size."""
        content = b"%PDF-1.4\n" + text.encode("utf-8") + b"\n%%EOF"
        iterations = 50
        start = time.perf_counter()
        for _ in range(iterations):
            detector.validate_content_integrity(content, "application/pdf")
        elapsed = time.perf_counter() - start
        avg_ms = (elapsed / iterations) * 1000
        # Validation should be fast regardless of content size
        assert avg_ms < 100, f"MIME validation too slow for {label}: {avg_ms:.1f}ms"


# ── Parser Benchmarks ──────────────────────────────────────────────

class TestParserBenchmarks:
    """Benchmark document parsing throughput and accuracy."""

    @pytest.fixture
    def parser(self):
        return ParserPipeline()

    @pytest.mark.parametrize("label,text", BENCHMARK_SIZES)
    @pytest.mark.asyncio
    async def test_txt_parse_throughput(self, parser, label, text):
        """Measure text parsing throughput at different sizes."""
        filename = f"bench_{label}.txt"
        path = TEST_FILES_DIR / filename
        path.write_text(text, encoding="utf-8")

        iterations = 5
        latencies = []
        for _ in range(iterations):
            start = time.perf_counter()
            result = await parser.parse(filename, str(path), "text/plain")
            latencies.append(time.perf_counter() - start)

        # Normalize line endings for comparison (Windows: \r\n from file write)
        expected_normalized = text.strip().replace('\r\n', '\n')
        actual_normalized = result.text.strip().replace('\r\n', '\n')
        assert actual_normalized == expected_normalized, \
            f"Text parsing should preserve content exactly for {label}"

        assert result.confidence == 1.0

    @pytest.mark.asyncio
    async def test_pdf_parse_accuracy(self, parser):
        """Verify PDF text extraction preserves key resume fields."""
        import fitz
        path = str(TEST_FILES_DIR / "bench_accuracy.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), RESUME_TEXT[:2000], fontsize=10)
        doc.save(path)
        doc.close()

        result = await parser.parse("bench_accuracy.pdf", path, "application/pdf")

        # Should preserve key information
        text = result.text
        assert "Software Engineer" in text
        assert "Acme Corp" in text or "acme" in text.lower()
        assert result.metadata["page_count"] == 1
        assert result.confidence > 0.85

    @pytest.mark.asyncio
    async def test_docx_parse_accuracy(self, parser):
        """Verify DOCX extraction preserves section structure."""
        from docx import Document
        path = str(TEST_FILES_DIR / "bench_accuracy.docx")
        doc = Document()
        doc.add_heading("RESUME", 0)
        doc.add_paragraph("John Doe - Software Engineer")
        doc.add_heading("EXPERIENCE", 1)
        doc.add_paragraph("Acme Corp, 2019-2023: Led engineering team of 10")
        doc.add_heading("SKILLS", 1)
        doc.add_paragraph("Python, React, TypeScript, AWS")
        doc.save(path)

        result = await parser.parse("bench_accuracy.docx", path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        text = result.text
        assert "John Doe" in text
        assert "Acme Corp" in text
        assert "Python" in text
        assert result.confidence > 0.9
        # Section headers should have been detected
        sections = result.metadata.get("sections", [])
        assert len(sections) >= 1


# ── OCR Benchmarks ─────────────────────────────────────────────────

class TestOcrBenchmarks:
    """Benchmark OCR accuracy and performance."""

    @pytest.fixture
    def ocr(self):
        return OcrPipeline()

    def _create_text_image(self, text: str, size=(800, 200)) -> bytes:
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 24)
        except (OSError, IOError):
            font = ImageFont.load_default()
        draw.text((20, 20), text, fill="black", font=font)

        import io
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()

    @pytest.mark.integration
    @requires_tesseract
    @pytest.mark.asyncio
    async def test_ocr_accuracy_known_text(self, ocr):
        """Measure OCR accuracy on a known text image."""
        expected = "John Doe\nSoftware Engineer\nPython Developer"
        img_bytes = self._create_text_image(expected)
        path = TEST_FILES_DIR / "ocr_bench_accuracy.png"
        path.write_bytes(img_bytes)

        result = await ocr.process("ocr_bench_accuracy.png", str(path), content_type="image/png")

        text_lower = result.text.lower()
        # Soft accuracy check — OCR on synthetic images varies by system
        words_in = sum(1 for word in ["john", "doe", "software", "engineer", "python"]
                       if word in text_lower)
        accuracy = words_in / 5
        assert accuracy >= 0.4, \
            f"OCR accuracy too low: {accuracy:.0%} matched. Text: {result.text[:100]}"

    @pytest.mark.integration
    @requires_tesseract
    @pytest.mark.asyncio
    async def test_ocr_performance_single_image(self, ocr):
        """OCR on a single image should complete within timeout."""
        img_bytes = self._create_text_image("Performance Test\nPython Engineer")
        path = TEST_FILES_DIR / "ocr_perf.png"
        path.write_bytes(img_bytes)

        iterations = 5
        latencies = []
        for _ in range(iterations):
            start = time.perf_counter()
            result = await ocr.process("ocr_perf.png", str(path), content_type="image/png")
            latencies.append(time.perf_counter() - start)

        avg_s = statistics.mean(latencies)
        # OCR is inherently slow (1-3s per image is acceptable)
        assert avg_s < 10.0, f"OCR too slow: {avg_s:.2f}s per image"

    @pytest.mark.integration
    @requires_tesseract
    def test_ocr_fallback_chain_executed(self, ocr):
        """The fallback chain should produce output even with difficult images."""
        # Create a very simple image that should be easy
        img_bytes = self._create_text_image("A B C D E F G")
        text, conf, method = ocr._ocr_image_bytes_with_fallback(img_bytes)
        assert isinstance(text, str) and len(text) > 0
        assert conf > 0.0

    @pytest.mark.integration
    @requires_tesseract
    def test_preprocessing_pipeline_all_methods_work(self, ocr):
        """Every preprocessing method should run without errors."""
        img_bytes = self._create_text_image("TEST IMAGE TEXT")
        for method in ("none", "grayscale", "sharpen", "contrast_boost", "binarize"):
            text, conf = ocr._ocr_image_bytes(img_bytes, preprocess=method)
            assert isinstance(text, str), f"Method {method} failed"


# ── End-to-end Pipeline Benchmark ──────────────────────────────────

class TestEndToEndBenchmark:
    """Full pipeline benchmarks from detection through extraction."""

    @pytest.fixture
    def detector(self):
        return MimeDetectorPipeline()

    @pytest.fixture
    def parser(self):
        return ParserPipeline()

    @pytest.fixture
    def ocr(self):
        return OcrPipeline()

    @pytest.mark.asyncio
    async def test_full_text_pipeline(self, detector, parser, ocr):
        """End-to-end: MIME → Parse → (skip OCR) for text files."""
        content = RESUME_TEXT.encode("utf-8")
        path = TEST_FILES_DIR / "e2e_bench.txt"
        path.write_bytes(content)

        # MIME detection
        mime = await detector.detect("e2e_bench.txt", content)
        assert mime == "text/plain"

        # Parsing
        parse_result = await parser.parse("e2e_bench.txt", str(path), mime)
        assert parse_result.text.strip() == RESUME_TEXT.strip()

        # OCR detection should skip
        needs_ocr = await ocr._detect_needs_ocr(
            "e2e_bench.txt", str(path),
            existing_text=parse_result.text,
        )
        assert needs_ocr is False

    @pytest.mark.integration
    @requires_tesseract
    @pytest.mark.asyncio
    async def test_full_image_pipeline(self, detector, parser, ocr):
        """End-to-end: MIME → Parse → OCR for image files."""
        from PIL import Image, ImageDraw, ImageFont
        import io

        img = Image.new("RGB", (800, 200), color="white")
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 24)
        except (OSError, IOError):
            font = ImageFont.load_default()
        draw.text((20, 20), "John Doe\nSoftware Engineer", fill="black", font=font)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        content = buf.getvalue()

        path = TEST_FILES_DIR / "e2e_image.png"
        path.write_bytes(content)

        # MIME detection
        mime = await detector.detect("e2e_image.png", content)
        assert mime == "image/png"

        # OCR should be triggered
        needs_ocr = await ocr._detect_needs_ocr(
            "e2e_image.png", str(path),
            existing_text="",
        )
        assert needs_ocr is True

        # Run OCR
        ocr_result = await ocr.process("e2e_image.png", str(path), content_type=mime)
        assert ocr_result.metadata["ocr_applied"] is True
        text_lower = ocr_result.text.lower()
        assert any(w in text_lower for w in ["john", "doe", "software"]), \
            f"OCR should extract text from image: {ocr_result.text[:100]}"


# ── Failure Recovery Benchmarks ────────────────────────────────────

class TestFailureRecovery:
    """Benchmark failure handling and error recovery paths."""

    @pytest.fixture
    def parser(self):
        return ParserPipeline()

    @pytest.fixture
    def ocr(self):
        return OcrPipeline()

    @pytest.mark.asyncio
    async def test_missing_file_gives_permanent_error(self, parser):
        """Missing file should produce PermanentPipelineError (not crash)."""
        with pytest.raises(PermanentPipelineError):
            await parser.parse("gone.pdf", "/tmp/nonexistent.pdf", "application/pdf")

    @pytest.mark.asyncio
    async def test_empty_file_gives_insufficient_error(self, parser):
        """Empty file should produce InsufficientContentError."""
        path = TEST_FILES_DIR / "empty_bench.txt"
        path.write_text("")
        with pytest.raises(InsufficientContentError):
            await parser.parse("empty_bench.txt", str(path), "text/plain")

    @pytest.mark.asyncio
    async def test_ocr_recovers_missing_file(self, ocr):
        """OCR on missing file should raise RetryablePipelineError."""
        with pytest.raises(RetryablePipelineError):
            await ocr.process("gone.png", "/tmp/gone.png", content_type="image/png")
