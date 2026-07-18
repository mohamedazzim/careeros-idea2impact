"""
Integration and unit tests for the production DocumentParser pipeline.
Tests PDF, DOCX, TXT, RTF, and image parsing with real extraction.
"""
import pytest
from pathlib import Path
from src.services.resume.processing.parser import (
    ParserPipeline,
    ParseResult,
    InsufficientContentError,
    PermanentPipelineError,
    RetryablePipelineError,
)

TEST_FILES_DIR = Path(__file__).parent / "test_files"


@pytest.fixture
def parser():
    return ParserPipeline()


@pytest.fixture(autouse=True)
def ensure_test_files():
    TEST_FILES_DIR.mkdir(exist_ok=True)


# ── Helpers ─────────────────────────────────────────────────────────

def _write_test_file(name: str, content: bytes) -> Path:
    path = TEST_FILES_DIR / name
    path.write_bytes(content)
    return path


# ── TXT Parsing ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_text_file(parser):
    content = b"John Doe\nSoftware Engineer\nSkills: Python, React, TypeScript\nExperience: 5 years at Acme Corp\nEducation: BS Computer Science"
    path = _write_test_file("resume.txt", content)
    result = await parser.parse("resume.txt", str(path), "text/plain")
    assert result.content_type == "text/plain"
    assert "John Doe" in result.text
    assert "Software Engineer" in result.text
    assert result.metadata["word_count"] >= 10
    assert result.confidence == 1.0


@pytest.mark.asyncio
async def test_parse_text_empty_content(parser):
    content = b"   \n   "
    path = _write_test_file("empty.txt", content)
    with pytest.raises(InsufficientContentError):
        await parser.parse("empty.txt", str(path), "text/plain")


@pytest.mark.asyncio
async def test_parse_text_insufficient_words(parser):
    content = b"hi"
    path = _write_test_file("short.txt", content)
    with pytest.raises(InsufficientContentError):
        await parser.parse("short.txt", str(path), "text/plain")


# ── Section Detection ──────────────────────────────────────────────

def test_detect_sections(parser):
    text = """SUMMARY
Experienced developer.

EXPERIENCE
Acme Corp, 2019-2023

EDUCATION
University of Technology, 2015-2019

SKILLS
Python, React, TypeScript"""
    sections = parser._detect_sections(text)
    assert "SUMMARY" in sections
    assert "EXPERIENCE" in sections
    assert "EDUCATION" in sections
    assert "SKILLS" in sections


def test_detect_sections_case_insensitive(parser):
    text = "Experience\nWork history\nSkills Summary"
    sections = parser._detect_sections(text)
    assert "Experience" in sections


# ── PDF Parsing ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_pdf_file(parser):
    """Integration test — parses a real PDF generated via PyMuPDF."""
    import fitz
    path = str(TEST_FILES_DIR / "test_resume.pdf")
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    page.insert_textbox(
        fitz.Rect(72, 72, 540, 720),
        "John Doe\nSoftware Engineer\n\nSkills: Python, React, TypeScript\nExperience: Acme Corp 2019-2023\nEducation: BS Computer Science",
        fontsize=11,
    )
    doc.save(path)
    doc.close()

    result = await parser.parse("test_resume.pdf", path, "application/pdf")
    assert result.content_type == "application/pdf"
    assert "John Doe" in result.text
    assert result.metadata["page_count"] == 1
    assert result.confidence > 0.7


@pytest.mark.asyncio
async def test_parse_pdf_multipage(parser):
    """Test multi-page PDF parsing."""
    import fitz
    path = str(TEST_FILES_DIR / "multipage.pdf")
    doc = fitz.open()
    for i in range(3):
        page = doc.new_page(width=612, height=792)
        page.insert_textbox(
            fitz.Rect(72, 72, 540, 720),
            f"Page {i + 1} Report\n\nContent for page {i + 1}\nAdditional details here for the report on page {i + 1}",
            fontsize=11,
        )
    doc.save(path)
    doc.close()

    result = await parser.parse("multipage.pdf", path, "application/pdf")
    assert result.metadata["page_count"] == 3
    assert "Page 1" in result.text
    assert "Page 2" in result.text
    assert "Page 3" in result.text


@pytest.mark.asyncio
async def test_parse_pdf_malformed(parser):
    """Malformed PDF should raise error."""
    path = _write_test_file("bad.pdf", b"%PDF-1.4\ncorrupted content\xff\xff\xff")
    with pytest.raises((RetryablePipelineError, InsufficientContentError, PermanentPipelineError)):
        await parser.parse("bad.pdf", str(path), "application/pdf")


@pytest.mark.asyncio
async def test_parse_pdf_file_not_found(parser):
    with pytest.raises(PermanentPipelineError):
        await parser.parse("missing.pdf", "/nonexistent/path.pdf", "application/pdf")


# ── DOCX Parsing ───────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_docx_file(parser):
    """Integration test — creates and parses a real DOCX."""
    from docx import Document
    path = str(TEST_FILES_DIR / "test_resume.docx")
    doc = Document()
    doc.add_heading("John Doe - Software Engineer", 0)
    doc.add_paragraph("Skills: Python, React, TypeScript")
    doc.add_heading("Experience", 1)
    doc.add_paragraph("Acme Corp, Senior Developer, 2019-2023")
    doc.add_heading("Education", 1)
    doc.add_paragraph("BS Computer Science, University of Technology, 2015")
    doc.save(path)
    
    result = await parser.parse("test_resume.docx", path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text_lower = result.text.lower()
    assert "john doe" in text_lower
    assert "software engineer" in text_lower
    assert "python" in text_lower
    assert "acme corp" in text_lower
    assert len(result.metadata.get("sections", [])) >= 1


# ── RTF Parsing ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_rtf_basic(parser):
    rtf_content = br'{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}\f0\fs24 John Doe, Software Engineer\par Skills: Python, React, TypeScript\par Experience: Acme Corp}'
    path = _write_test_file("resume.rtf", rtf_content)
    result = await parser.parse("resume.rtf", str(path), "application/rtf")
    text_lower = result.text.lower()
    has_text = any(w in text_lower for w in ["john", "doe", "software", "python", "acme"])
    assert has_text, f"Expected recognizable text, got: {result.text}"


# ── Metadata output ────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_result_has_metadata(parser):
    content = b"John Doe\nSoftware Engineer\nSkills: Python, React, TypeScript\nExperience: 5 years at Acme\nEducation: BS Computer Science"
    path = _write_test_file("meta.txt", content)
    result = await parser.parse("meta.txt", str(path), "text/plain")
    assert result.metadata is not None
    assert "word_count" in result.metadata
    assert "sections" in result.metadata
    assert "parse_duration_ms" in result.metadata
    assert result.metadata["word_count"] > 5


# ── LangGraph node interface ───────────────────────────────────────

@pytest.mark.asyncio
async def test_parser_langgraph_node_success(parser):
    content = b"John Doe - Senior Engineer\nSkills: Python, React, TypeScript\nExperience: 5 years"
    path = _write_test_file("lg_test.txt", content)
    state = {
        "filename": "lg_test.txt",
        "storage_path": str(path),
        "content_type": "text/plain",
    }
    update = await parser.run(state)
    assert update["raw_text"] is not None
    assert update["parse_error"] is None
    assert "John Doe" in update["raw_text"]


@pytest.mark.asyncio
async def test_parser_langgraph_node_error(parser):
    state = {
        "filename": "missing.txt",
        "storage_path": "/does/not/exist.txt",
        "content_type": "text/plain",
    }
    update = await parser.run(state)
    assert update["raw_text"] is None
    assert update["parse_error"] is not None
